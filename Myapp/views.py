from django.shortcuts import render

# Create your views here.
from django.http import HttpResponse,HttpResponseRedirect
import datetime
from Myapp.models import *
import os,shutil,json
import subprocess
import threading
import time
import re
import zipfile
from django.core.paginator import Paginator,EmptyPage,PageNotAnInteger
from django.contrib.auth.decorators import login_required
from django.contrib import auth
from docx import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from django.contrib.auth.models import User

# 进入登录页
def index(request):
    next = request.GET.get('next',None)
    return render(request,'index.html',{"next":next})

# 进入登录页
def index_error(request):
    return render(request,'index.html',{"error_msg":"密码错误！"})

# 改密码
def changePWD(request):
    username = request.GET['username']
    new_pwd = request.GET['new_pwd']
    user = User.objects.get(username=username)
    user.set_password(new_pwd)
    user.save()
    return HttpResponse('修改成功')

# 登录功能
def login_action(request):
    next = request.GET.get('next','/home/')
    if next == '':
        next = '/home/'
    username = request.POST.get('Username',None)
    password = request.POST.get('Password',None)
    user = auth.authenticate(username=username,password=password)
    if user is not None:
        auth.login(request,user)
        request.session['user'] = username
        return HttpResponseRedirect(next)
    else:
        return HttpResponseRedirect('/index_error/')

# 注册功能
def register_action(request):
    username = request.POST.get('Username', None)
    password = request.POST.get('Password', None)
    try:
        user = User.objects.create_user(username=username,password=password)
        user.save()
        auth.login(request,user)
        request.session['user'] = username
        return HttpResponseRedirect('/home/')
    except:
        return HttpResponse('注册失败！用户名已经存在！')

# 退出功能
def logout(request):
    auth.logout(request)
    return HttpResponseRedirect('/index/')


# 进入首页
@login_required
def home(request):
    res = {}
    res["username"] = request.user.username
    res["time"] = datetime.datetime.strftime(datetime.datetime.now(), '%c')
    res["href"] = DB_href.objects.all()
    res['duan'] = DB_duan.objects.all()
    return render(request,'home.html',res)

# 新增端
def save_duan(request):
    duan_name = request.GET['duan_name']
    duan_host = request.GET['duan_host']
    duan = DB_duan.objects.create(name=duan_name,host=duan_host,user=request.user.username)
    # 复制出自己的调试包
    new_client_name = 'client_'+str(duan.id)
    demo_path = 'MyClient/demo_client'
    new_path = 'MyClient/'+new_client_name
    shutil.copytree(demo_path,new_path)
    return HttpResponse()

# 删除端
def del_duan(request,did):
    duan = DB_duan.objects.filter(id=did)[0]
    if request.user.username == duan.user:
        duan.delete()
    else: # 查看后门权限
        users = DB_quanxian.objects.filter(name='删除端')[0].users
        if request.user.username in users.split(','):
            duan.delete()
    return HttpResponse('')

# 进入用例列表页
@login_required
def case_list(request,did):
    kwords = request.GET.get("kwords",'')

    cases = DB_case.objects.filter(duan_id=did,name__contains=kwords)

    p = Paginator(cases,5,1)
    page = request.GET.get('page')
    try:
        new_cases = p.page(page)
    except PageNotAnInteger:
        new_cases = p.page(1)
    except EmptyPage:
        new_cases = p.page(p.num_pages)

    res = {}
    res['cases'] = new_cases
    res['duan'] = DB_duan.objects.filter(id = did)[0]
    res["href"] = DB_href.objects.all()
    res["hosts"] =  DB_duan.objects.filter(id = did)[0].host.split(',')
    res["username"] = request.user.username
    return render(request,'case.html',res)

# 添加用例
def add_case(request,did):
    DB_case.objects.create(duan_id=did)
    return HttpResponseRedirect('/case_list/'+did+'/')
    # return case_list(request,did)

# 删除用例
def del_case(request,cid):
    case = DB_case.objects.filter(id = cid)
    duan_id = case[0].duan_id
    duan = DB_duan.objects.filter(id=duan_id)[0]
    if request.user.username == duan.user:
        case.delete()
    else: # 查看后门权限
        users = DB_quanxian.objects.filter(name='删除用例')[0].users
        if request.user.username in users.split(','):
            case.delete()
    return HttpResponseRedirect('/case_list/' + duan_id + '/')

# 获取用例设置的数据
def set_case(request):
    case_id = request.GET['id']
    case = list( DB_case.objects.filter(id=case_id).values() )[0]
    return HttpResponse(json.dumps(case),content_type="application/json")

# 保存用例
def save_case(request):
    case_id = request.GET['case_id']
    case_name = request.GET['case_name']
    case_counts = request.GET['case_counts']
    BF = request.GET['BF']
    JK = request.GET['JK']
    DB_case.objects.filter(id=case_id).update(name=case_name,counts=case_counts,BF=BF,JK=JK)
    return HttpResponse('')

# 保存监控设置
def save_monitor(request):
    monitor_host = request.GET['monitor_host']
    monitor_time = request.GET['monitor_time']
    monitor_phone = request.GET['monitor_phone']
    monitor_email = request.GET['monitor_email']
    monitor_dingtalk = request.GET['monitor_dingtalk']
    duan_id = request.GET['duan_id']

    DB_duan.objects.filter(id=duan_id).update(monitor_host=monitor_host,monitor_time=monitor_time,
                                              monitor_phone=monitor_phone,
                                              monitor_email=monitor_email,
                                              monitor_dingtalk=monitor_dingtalk)

    return HttpResponse('')

# 保存旧端
def save_old_duan(request):
    name = request.GET['name']
    host = request.GET['host']
    user = request.GET['user']
    max_bf = request.GET['max_bf']
    duan_id = request.GET['duan_id']

    DB_duan.objects.filter(id=duan_id).update(name=name,host=host,user=user,max_bf=max_bf)

    return HttpResponse('')

# 上传脚本
def upload_py(request,cid):
    # 拿到端id
    duan_id = DB_case.objects.filter(id=cid)[0].duan_id
    # 获取到上传的脚本
    myFile = request.FILES.get("fileUpload",None)
    if not myFile:
        return HttpResponseRedirect('/case_list/%s/'%duan_id)
    file_name = str(myFile)
    fp = open('MyClient/client_%s/cases/%s'%(duan_id,file_name),'wb+')
    for i in myFile.chunks():
        fp.write(i)
    fp.close()
    # 更新用例的数据库py字端
    DB_case.objects.filter(id=cid).update(py=file_name)
    # 返回
    return HttpResponseRedirect('/case_list/%s/'%duan_id)

# 执行脚本
def run_case(request):
    case_id = request.GET['case_id']
    host = request.GET['host']
    case = DB_case.objects.filter(id=case_id)[0]
    conuts = case.counts
    duan_id = case.duan_id
    py = case.py #即是py脚本，也可能excel表后缀.xls
    if py in ['',None,' ','None']:
        return HttpResponse('Error')
    duan = DB_duan.objects.filter(id=duan_id)[0]
    if request.user.username == duan.user:
        if '.py' in py:
            subprocess.call('python3 MyClient/client_%s/cases/%s %s %s'%(duan_id,py,host,conuts),shell=True)
        elif '.xls' in py:
            subprocess.call('python3 MyClient/client_%s/public/make_script.py %s %s %s'%(duan_id,py,host,conuts),shell=True)
    return HttpResponse('')

def waibu_run_case(request):
    case_id = request.GET['case_id']
    host = request.GET['host']
    case = DB_case.objects.filter(id=case_id)[0]
    duan_id = case.duan_id
    py = case.py
    if py in ['', None, ' ', 'None']:
        return HttpResponse('Error')
    subprocess.call('python3 MyClient/client_%s/cases/%s %s' % (duan_id, py, host), shell=True)
    res = {"errcode":200,"errmsg":"","content":{}}
    return HttpResponse(json.dumps(res),content_type='application/json')

# 并发用例
def bf_case(request,did):
    # 所有需要并发的case先拿出来
    duan = DB_duan.objects.filter(id=did)[0]
    if request.user.username != duan.user:
        return HttpResponse('权限不足')

    cases = DB_case.objects.filter(duan_id=did,BF=True)
    max_bf = DB_duan.objects.filter(id=did)[0].max_bf
    host = request.GET['host']
    def do_case(case): # 线程要运行的函数
        if '.py' in case.py:
            subprocess.call('python3 MyClient/client_%s/cases/%s %s %s' % (case.duan_id, case.py,host,case.counts), shell=True)
        elif '.xls' in case.py:
            subprocess.call('python3 MyClient/client_%s/public/make_script.py %s %s %s'%(case.duan_id,case.py,host,case.conuts),shell=True)
        print('执行完：',case.name)

    ts = [] # 声明空的线程池
    for case in cases:
        if case.py not in  ['',None,' ','None']:
            t =threading.Thread(target=do_case,args=(case,)) # 声明子线程
            t.setDaemon(True) # 设置守护线程
            ts.append(t) # 添加到线程池

    for i in range(0,len(ts),max_bf):
        tmp = ts[i:i+max_bf]
        for t in tmp: # 启动线程
            t.start()
        for t in tmp: # 让主线程等待子线程
            t.join()

    time.sleep(2)
    print('全部执行完毕')
    return look_reports(request,did=did)

# 启动监控
def start_monitor(request,did):
    # 先判断监控线程是否已经启动
    duan = DB_duan.objects.filter(id=did)[0]
    if request.user.username != duan.user:
        return HttpResponse('权限不足')

    try:
        subprocess.check_output('ps -ef | grep "monitor.py %s WEB" | grep -v "grep"'%did,shell=True)
        print('监控已经启动')
        return HttpResponseRedirect('/case_list/'+did+'/')
    except:
        print('监控无，可以启动')
    # 如果没启动，则启动
    def start_server():
        subprocess.call('python3 Myapp/monitor.py %s WEB'%did,shell=True)
    t = threading.Thread(target=start_server)
    t.setDaemon(True)
    t.start()
    # 返回
    return HttpResponseRedirect('/case_list/'+did+'/')

# 关闭监控
def stop_monitor(request,did):
    duan = DB_duan.objects.filter(id=did)[0]
    if request.user.username != duan.user:
        return HttpResponse('权限不足')
    # 找到监控线程
    try:
        cop = subprocess.check_output('ps -ef | grep "monitor.py %s WEB" | grep -v "grep"' % did, shell=True)
        print('监控已经启动，即将停止')
        ports = re.findall(r'(\d+)',str(cop))
        max_id = max([int(i) for i in ports])
        subprocess.call('kill -9 %s'%str(max_id),shell=True)
        print('监控服务停止！')
    except:
        print('监控无，无需停止')
    # 返回
    return HttpResponseRedirect('/case_list/'+did+'/')

# 查看报告
def look_report(request,cid):
    case = DB_case.objects.filter(id=cid)[0]
    duan_id = case.duan_id
    py = case.py
    try:
        report_name = py.replace('.py','.html').replace('.xls','.html')
    except:
        return HttpResponse('该用例尚未上传脚本！')
    try:
        return render(request,'client_%s/reports/'%duan_id+report_name)
    except:
        return HttpResponse('该用例报告不存在！')

# 压缩函数
def get_zip_file(input_path, result):
    """
    对目录进行深度优先遍历
    :param input_path:
    :param result:
    :return:
    """
    files = os.listdir(input_path)
    for file in files:
        if os.path.isdir(input_path + '/' + file):
            get_zip_file(input_path + '/' + file, result)
        else:
            result.append(input_path + '/' + file)
def zip_file_path(input_path, output_path, output_name):
    """
    压缩文件
    :param input_path: 压缩的文件夹路径
    :param output_path: 解压（输出）的路径
    :param output_name: 压缩包名称
    :return:
    """
    f = zipfile.ZipFile(output_path + '/' + output_name, 'w', zipfile.ZIP_DEFLATED)
    filelists = []
    get_zip_file(input_path, filelists)
    for file in filelists:
        f.write(file)
    # 调用了close方法才会保证完成压缩
    f.close()

# 下载本地调试包
def download_client(request,did):
    zip_name = 'MyClient/CLIENT_%s.zip'%did
    # 删除旧的压缩包
    try:os.remove(zip_name)
    except:pass
    # 压缩client包
    zip_file_path('MyClient/client_%s'%did,'MyClient','CLIENT_%s.zip'%did)
    time.sleep(1)

    # 返回给前端这个压缩包
    try:
        file = open(zip_name,'rb')
    except:
        return HttpResponse('')

    response = HttpResponse(file)
    response['Content-Type'] = 'application/octet-stream'
    response['Content-Disposition'] = 'attachment;filename="%s"'%'CLIENT_%s.zip'%did
    time.sleep(1)
    # 删除新的压缩包
    try:
        os.remove(zip_name)
    except:
        pass
    return response

# 查看报告总结
def look_reports(request,did=''):
    try:
        duan_id = request.GET['duan_id']
    except:
        duan_id = did

    pys =  list(DB_case.objects.filter(duan_id=duan_id,BF=True).values())
    # 提取报告内容
    res = '【并发结果如下】\n'
    all_case = 0
    pass_case = 0
    errfail_list = []
    for i in pys:
        errfail_case = 0
        report_name = 'MyClient/client_%s/reports/%s'%(duan_id,str(i['py']).replace('.py','.html').replace('.xls','.html'))
        try:
            # 解析报告
            with open(report_name,'r') as fp:
                text = fp.read()
                res_data = re.findall(r'<td name>(.*?)</td>',text)
                all_case += int(res_data[0])
                pass_case += int(res_data[1])
                errfail_case += int(res_data[2])+int(res_data[3])
                if errfail_case > 0:
                    errfail_list.append(i['name'])
        except Exception as e:
            print(e)
    res += '小用例总数:%s \n通过的小用例总数:%s \n失败和报错的小用例总数:%s \n'%(str(all_case),str(pass_case),str(errfail_case))
    res += '出错或失败的用例为:'+','.join(errfail_list)
    res += '\n具体报告详情请点击用例后的查看按钮'
    return HttpResponse(res)

# 下载导出测试报告
def download_reports(request,did):
    httpresponse = look_reports(request,did=did)
    Duan = DB_duan.objects.filter(id=did)[0]

    # 创建word文档
    doc = Document()
    doc.styles['Normal'].font.name = '微软雅黑'

    # 写标题
    p = doc.add_paragraph('【%s 的自动化测试报告】'%Duan.name)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中

    p = doc.add_paragraph('')
    p.add_run('本测试报告由xxx的web自动化平台:xxxxxx生成,地址:xxxxxx').font.color.rgb = RGBColor(*(172,182,182))
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中

    p = doc.add_paragraph('')
    p.add_run('报告生成时间: %s'%time.strftime('%Y-%m-%d %H:%M:%s')).font.color.rgb = RGBColor(*(172,182,182))
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中

    doc.add_paragraph('报告人: %s'%request.user.username,style='Intense Quote')
    doc.add_paragraph('平台维护: %s'%'测试开发干货',style='Intense Quote')
    doc.add_paragraph('执行环境: %s'%'http://', style='Intense Quote')

    doc.add_paragraph('用例执行结果总览如下:',style='Intense Quote')

    p = doc.add_paragraph(httpresponse.content.decode())
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('用例执行结果传送门如下:',style='Intense Quote')

    # 获取所有报告的详情地址并展示到word中
    cases = DB_case.objects.filter(duan_id=did,BF=True)
    for i in cases:
        doc.add_paragraph('用例名字:'+str(i.name))
        doc.add_paragraph('脚本名字:'+str(i.py))
        doc.add_paragraph('报告地址:' + 'http://xxxx:8000/look_report/'+str(i.id)+'/')
        doc.add_paragraph(' ',style='Intense Quote')

    # 保存并下载
    file_name  = 'TESTREPORTS_%s.docx'%time.strftime('%Y-%m-%d %H:%M:%s')
    doc.save(file_name)
    with open(file_name,'rb') as fp:
        response = HttpResponse(fp)
        response['Content-Type'] = 'application/octet-stream'
        response['Content-Disposition'] = 'attachment;filename="%s"'%file_name

    os.remove(file_name)

    return response

# 上传tools.py
def upload_tools(request,did):
    duan = DB_duan.objects.filter(id=did)[0]
    if request.user.username != duan.user:
        return HttpResponse('权限不足')

    myFile = request.FILES.get("file_tools",None)
    if not myFile:
        return HttpResponseRedirect("/case_list/"+did)

    filename = str(myFile)
    fp = open(os.path.join("MyClient/client_%s/public/"%did,filename),'wb+')
    for chunk in myFile.chunks():
        fp.write(chunk)
    fp.close()

    return HttpResponseRedirect("/case_list/"+did)


# 打开元素库
def object_in(request,did):
    objects = DB_object.objects.filter(duan_id=did)
    res = {}
    res['objects'] = objects
    res['duan'] = DB_duan.objects.filter(id=did)[0]
    return render(request,'object.html',res)

# 新增元素
def add_object(request,did):
    DB_object.objects.create(duan_id=did)
    return HttpResponseRedirect('/object_in/'+did+'/')

# 保存元素
def save_object(request):
    id = request.GET['id']
    name = request.GET['name']
    page = request.GET['page']
    tmp_method = request.GET['tmp_method']
    tmp_value = request.GET['tmp_value']
    tag = request.GET['tag']
    index = request.GET['index']
    DB_object.objects.filter(id=id).update(name=name,page=page,tmp_method=tmp_method,tmp_value=tmp_value,tag=tag,index=int(index))
    return HttpResponse('')

# 删除元素
def del_object(request,oid):
    did  = DB_object.objects.filter(id=oid)[0].duan_id
    DB_object.objects.filter(id=oid).delete()
    return HttpResponseRedirect('/object_in/'+did+'/')

# 外部获取元素信息
def waibu_get_object(request,oid):
    object = list(DB_object.objects.filter(id=oid).values())[0]
    return HttpResponse(json.dumps(object),content_type='application/json')

# 外部更新元素信息
def waibu_update_object(request,oid):
    tmp_method = request.POST.get('tmp_method')
    tmp_value = request.POST.get('tmp_value')
    index = request.POST.get('index')
    tag = request.POST.get('tag')
    DB_object.objects.filter(id=oid).update(tmp_method=tmp_method,tmp_value=tmp_value,index=index,tag=tag)
    return HttpResponse("")